"""知识文档解析与分块"""

from __future__ import annotations

import fnmatch
import hashlib
import re
import uuid
from dataclasses import dataclass
from pathlib import Path

from config.settings import settings
from src.rag.models import TextChunk

SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}
CHUNK_SIZE = 1000
STEP_CHUNK_SIZE = 1200
CHUNK_OVERLAP = 150

REGION_LABELS = {"hk": "香港", "cn": "国内", "": "通用"}

SECTION_START_RE = re.compile(
    r"^(?:"
    r"[一二三四五六七八九十百千]+、|"
    r"（[一二三四五六七八九十]+）|"
    r"\d+\."  # 国内流程步骤，如 1.进群打招呼
    r")"
)


@dataclass
class StepSection:
    region: str
    step_title: str
    step_id: str
    body: str


def file_content_hash(path: Path) -> str:
    data = path.read_bytes()
    return hashlib.sha256(data).hexdigest()


def load_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        parts = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(parts)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            rows: list[str] = []
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                parts.append("[表格]\n" + "\n".join(rows))
        return "\n\n".join(parts)
    raise ValueError(f"不支持的文件类型: {path.suffix}")


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def is_section_start(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped in ("香港", "香港：") or stripped.startswith("香港："):
        return True
    if stripped == "国内":
        return True
    if re.match(r"^#{1,6}\s", stripped):
        return True
    return bool(SECTION_START_RE.match(stripped))


def looks_like_step_doc(text: str) -> bool:
    return bool(re.search(r"^[一二三四五六七八九十]+、", text, re.MULTILINE)) or "国内" in text


def split_into_step_sections(text: str) -> list[StepSection]:
    text = normalize_text(text)
    if not text:
        return []

    lines = text.split("\n")
    current_region = "hk"
    step_counter = {"hk": 0, "cn": 0}
    sections: list[StepSection] = []
    current_lines: list[str] = []
    current_title = ""

    def flush() -> None:
        nonlocal current_lines, current_title
        if not current_lines:
            return
        body = "\n".join(current_lines).strip()
        if not body:
            current_lines = []
            return
        step_counter[current_region] += 1
        title = current_title or body.split("\n", 1)[0][:80]
        sections.append(
            StepSection(
                region=current_region,
                step_title=title,
                step_id=f"{current_region}_{step_counter[current_region]}",
                body=body,
            )
        )
        current_lines = []

    for line in lines:
        stripped = line.strip()
        if stripped in ("香港", "香港：") or stripped.startswith("香港："):
            flush()
            current_region = "hk"
            current_title = "香港"
            current_lines = []
            continue
        if stripped == "国内":
            flush()
            current_region = "cn"
            current_title = "国内"
            current_lines = []
            continue
        if is_section_start(line):
            flush()
            current_title = stripped
            current_lines = [line]
        else:
            if not current_lines and stripped:
                if not current_title:
                    current_title = stripped[:80]
            current_lines.append(line)

    flush()
    return sections


CAUTION_SPLIT_RE = re.compile(r"(?<=\n)(?=注意事项[:：])")


def _split_body_at_caution(body: str) -> list[tuple[str, str]]:
    """Split section body into script and caution parts at 注意事项：."""
    parts = CAUTION_SPLIT_RE.split(body, maxsplit=1)
    if len(parts) == 1:
        return [(parts[0].strip(), "script")] if parts[0].strip() else []
    result: list[tuple[str, str]] = []
    if parts[0].strip():
        result.append((parts[0].strip(), "script"))
    if parts[1].strip():
        result.append((parts[1].strip(), "caution"))
    return result


def _caution_step_title(section: StepSection, caution_body: str) -> str:
    for line in caution_body.split("\n"):
        stripped = line.strip()
        if not stripped or stripped in ("注意事项", "注意事项："):
            continue
        if stripped.startswith("面签注意事项"):
            return "面签注意事项"
        if stripped.endswith("注意事项") and len(stripped) <= 40:
            return stripped.rstrip("：:")
    short = section.step_title.replace("\n", " ")[:20]
    return f"{short}·注意事项"


def _split_long_section(text: str, max_size: int = STEP_CHUNK_SIZE) -> list[str]:
    if len(text) <= max_size:
        return [text]

    parts = re.split(r"(?<=\n)(?=注意事项[:：])", text)
    chunks: list[str] = []
    buf = ""
    for part in parts:
        candidate = f"{buf}\n{part}".strip() if buf else part.strip()
        if len(candidate) <= max_size:
            buf = candidate
            continue
        if buf:
            chunks.append(buf)
        if len(part) <= max_size:
            buf = part.strip()
        else:
            start = 0
            while start < len(part):
                end = min(start + max_size, len(part))
                piece = part[start:end].strip()
                if piece:
                    chunks.append(piece)
                if end >= len(part):
                    break
                start = max(end - CHUNK_OVERLAP, start + 1)
            buf = ""
    if buf:
        chunks.append(buf)
    return chunks


def sections_to_chunk_specs(sections: list[StepSection]) -> list[tuple[str, dict[str, str]]]:
    specs: list[tuple[str, dict[str, str]]] = []
    for section in sections:
        region_label = REGION_LABELS.get(section.region, section.region)
        for body_part, chunk_kind in _split_body_at_caution(section.body):
            if chunk_kind == "caution":
                step_title = _caution_step_title(section, body_part)
                short = step_title[:30]
            else:
                step_title = section.step_title
                short = section.step_title.replace("\n", " ")[:30]
            prefix = f"[{region_label}·{short}] "
            full_body = f"{prefix}{body_part}"
            for piece in _split_long_section(full_body, STEP_CHUNK_SIZE):
                meta = {
                    "region": section.region,
                    "step_title": step_title,
                    "step_id": section.step_id,
                    "chunk_kind": chunk_kind,
                }
                specs.append((piece, meta))
    return specs


def split_into_chunks_legacy(text: str) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []

    sections: list[str] = []
    current: list[str] = []
    for line in text.split("\n"):
        if re.match(r"^#{1,6}\s", line) and current:
            sections.append("\n".join(current).strip())
            current = [line]
        else:
            current.append(line)
    if current:
        sections.append("\n".join(current).strip())

    chunks: list[str] = []
    for section in sections:
        if len(section) <= CHUNK_SIZE:
            if section:
                chunks.append(section)
            continue
        start = 0
        while start < len(section):
            end = min(start + CHUNK_SIZE, len(section))
            piece = section[start:end].strip()
            if piece:
                chunks.append(piece)
            if end >= len(section):
                break
            start = max(end - CHUNK_OVERLAP, start + 1)
    return chunks


def split_into_chunks(text: str) -> list[tuple[str, dict[str, str]]]:
    if looks_like_step_doc(text):
        sections = split_into_step_sections(text)
        if sections:
            return sections_to_chunk_specs(sections)
    legacy = split_into_chunks_legacy(text)
    return [(piece, {}) for piece in legacy]


def parse_document(path: Path, *, doc_id: str | None = None) -> tuple[str, list[TextChunk]]:
    path = path.resolve()
    doc_id = doc_id or str(uuid.uuid4())
    title = path.stem
    raw = load_text(path)
    piece_specs = split_into_chunks(raw)
    rel_path = _relative_source_path(path)
    chunks: list[TextChunk] = []
    for idx, (piece, meta) in enumerate(piece_specs):
        chunks.append(
            TextChunk(
                chunk_id=str(uuid.uuid4()),
                doc_id=doc_id,
                chunk_index=idx,
                text=piece,
                source_path=rel_path,
                title=title,
                token_count=len(piece),
                region=str(meta.get("region") or ""),
                step_title=str(meta.get("step_title") or ""),
                step_id=str(meta.get("step_id") or ""),
                chunk_kind=str(meta.get("chunk_kind") or "script"),
            )
        )
    return doc_id, chunks


def _relative_source_path(path: Path) -> str:
    from config.settings import PROJECT_ROOT

    try:
        return str(path.resolve().relative_to(PROJECT_ROOT)).replace("\\", "/")
    except ValueError:
        return str(path).replace("\\", "/")


def should_exclude_path(path: Path) -> bool:
    name = path.name
    if name.startswith("~$"):
        return True
    rel = _relative_source_path(path)
    for pattern in settings.rag_exclude_pattern_list():
        if fnmatch.fnmatch(name, pattern) or fnmatch.fnmatch(rel, pattern):
            return True
    return False


def iter_knowledge_files(directory: Path) -> list[Path]:
    if not directory.is_dir():
        return []
    files: list[Path] = []
    for path in sorted(directory.rglob("*")):
        if not path.is_file():
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            continue
        if should_exclude_path(path):
            continue
        files.append(path)
    return files


def chunk_stats(chunks: list[TextChunk]) -> dict[str, int]:
    stats: dict[str, int] = {}
    for c in chunks:
        key = c.region or "unknown"
        stats[key] = stats.get(key, 0) + 1
    return stats
